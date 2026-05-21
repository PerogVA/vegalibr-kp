# -*- coding: utf-8 -*-
"""ВЕГАЛИБР — веб-калькулятор КП"""

import os
import re
from io import BytesIO
from flask import (Flask, render_template, request, session,
                   redirect, url_for, send_file, jsonify)
from pricing import calc_item, calc_item_from_excel, parse_caliber, _PLUG_COMBO, NDS
from excel_parser import parse_excel
from text_parser import parse_pdf, parse_text
from doc_builder import build_kp
from stock import lookup as stock_lookup
# drawing_parser импортируется внутри маршрутов (lazy import)

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "vegalibr-secret-2025")

APP_PASSWORD = os.environ.get("APP_PASSWORD", "vegalibr2025")

# ─────────────────────────────────────────────────────────────────────────────
@app.route("/")
def index():
    if not session.get("auth"):
        return redirect(url_for("login"))
    return render_template("index.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        if request.form.get("password") == APP_PASSWORD:
            session["auth"] = True
            return redirect(url_for("index"))
        error = "Неверный пароль"
    return render_template("login.html", error=error)

@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))

# ─────────────────────────────────────────────────────────────────────────────
@app.route("/parse", methods=["POST"])
def parse():
    """Парсит Excel-файл, возвращает JSON со списком позиций."""
    if not session.get("auth"):
        return jsonify({"error": "not_auth"}), 403

    include_kalib = request.form.get("include_kalib", "true").lower() == "true"
    try:
        discount = float(request.form.get("discount", "20")) / 100.0
        discount = max(0.0, min(0.99, discount))
    except ValueError:
        discount = 0.20

    # Текст из буфера обмена?
    raw_text = request.form.get("raw_text", "").strip()
    if raw_text:
        items_raw, err = parse_text(raw_text)
        if err:
            return jsonify({"error": err}), 400
    else:
        f = request.files.get("file")
        if not f:
            return jsonify({"error": "Загрузите файл или вставьте текст"}), 400

        fname = (f.filename or '').lower()
        file_buf = BytesIO(f.read())
        try:
            if fname.endswith('.pdf'):
                items_raw, err = parse_pdf(file_buf)
            elif fname.endswith('.docx'):
                # Извлекаем текст из Word-файла
                from docx import Document as _DocxDoc
                _d = _DocxDoc(file_buf)
                _lines = [p.text for p in _d.paragraphs]
                for _t in _d.tables:
                    for _r in _t.rows:
                        _lines.append('\t'.join(c.text for c in _r.cells))
                items_raw, err = parse_text('\n'.join(_lines))
            elif any(fname.endswith(ext) for ext in ('.png', '.jpg', '.jpeg', '.bmp', '.tif', '.tiff')):
                return jsonify({"error": "Перетащи изображение в поле загрузки и нажми «📐 Анализировать чертёж»"}), 400
            else:
                items_raw, err = parse_excel(file_buf)
        except Exception as e:
            return jsonify({"error": f"Не удалось открыть файл: {e}"}), 400
        if err:
            return jsonify({"error": err}), 400

    result = []
    aa = []
    for name, qty, price_excel, kalib_excel in items_raw:
        # 1) Пробуем наш прайс (с разбивкой ПР/НЕ если нужно)
        rows = _expand(name, qty, include_kalib, discount)
        if rows:
            result.extend(rows)
            continue

        # 2) Если в Excel есть цена — используем её (без разбивки)
        if price_excel is not None:
            kal = (kalib_excel or 0.0) if include_kalib else 0.0
            row = calc_item_from_excel(name, qty, price_excel, kal,
                                       include_kalib=include_kalib, discount=discount)
            row['source'] = 'excel'
            _add_stock(row)
            result.append(row)
            continue

        # 3) Не удалось — в АА
        aa.append({"name": name, "qty": qty})

    return jsonify({"items": result, "aa": aa})


def _extract_qty_from_name(name):
    """Извлекает кол-во, вписанное в наименование (напр. 'ПР12шт' → 12). None если нет."""
    m = re.search(r'[-×х]?\s*(\d+)\s*шт\.?', name, flags=re.IGNORECASE)
    return int(m.group(1)) if m else None


def _infer_caliber_type(name):
    """
    Если имя начинается с 'Калибр М...' (без кольцо/пробка) или просто 'М...',
    пытается определить тип по квалитету:
      H/Н (прописные) → пробка (отверстие)
      g/e/f/r/d (строчные) → кольцо (вал)
    Возвращает нормализованное имя или исходное если определить не удалось.
    """
    # Голое «М4×0,7-6g» (без типа калибра) — инферим тип и добавляем префикс
    if re.match(r'^[МмMm]\s*\d', name):
        nl = name.lower()
        if 'кольцо' in nl or 'пробка' in nl or 'калибр' in nl:
            return name  # уже есть тип
        m = re.search(
            r'[МмMm]\s*\d+[,.]?\d*\s*[×xхХ]?\s*\d*[,.]?\d*\s*-\s*([A-Za-zА-Яа-яЁё0-9]+)',
            name)
        if m:
            qual = m.group(1)
            if re.search(r'[gefdr]', qual):
                return 'Калибр-кольцо ' + name
            if re.search(r'[HН]', qual):
                return 'Калибр-пробка ' + name
        return name

    # Работаем только с «Калибр М...»
    if not re.match(r'^Калибр\s+[МмMm]\s*\d', name, flags=re.IGNORECASE):
        return name
    # Уже есть тип — не трогаем
    nl = name.lower()
    if 'кольцо' in nl or 'пробка' in nl:
        return name

    # Ищем квалитет после дефиса: -5Н6Н, -6g, -6eR и т.п.
    m = re.search(r'[МмMm]\s*\d+[,.]?\d*\s*[×xхХ]?\s*\d*[,.]?\d*\s*-\s*([A-Za-zА-Яа-яЁё0-9]+)', name)
    if m:
        qual = m.group(1)
        # Если есть строчные латинские (g, e, f, d, r) → кольцо (вал)
        if re.search(r'[gefdr]', qual):
            suffix = name[len('калибр'):].strip() if name.lower().startswith('калибр') else name
            return 'Калибр-кольцо ' + suffix
        # Если есть прописные H/Н → пробка (отверстие)
        if re.search(r'[HН]', qual):
            suffix = name[len('калибр'):].strip() if name.lower().startswith('калибр') else name
            return 'Калибр-пробка ' + suffix
    return name


def _clean_name(name):
    """Убирает служебные аннотации и нормализует наименование к стандарту ВЕГАЛИБР."""
    # ГОСТ ссылки: ГОСТ 17758-72
    name = re.sub(r'\s*ГОСТ\s*\d+[\.\-–]\d+', '', name, flags=re.IGNORECASE)
    # Артикулы: 8221-3013, 8221-3013-6Н и подобные
    name = re.sub(r'\b\d{3,}[\-–]\d{3,}[\-–]?\w*\b', '', name)
    # >50мм / > 50 мм / более 50мм
    name = re.sub(r'\s*>?\s*50\s*мм\b', '', name, flags=re.IGNORECASE)
    name = re.sub(r'\s*более\s*\d+\s*мм?\b', '', name, flags=re.IGNORECASE)
    # Кол-во в имени: ×12шт / 12шт / 12 шт. / -12шт
    name = re.sub(r'\s*[-×х]?\s*\d+\s*шт\.?', '', name, flags=re.IGNORECASE)
    # Нормализация типа
    name = re.sub(r'^Кольцо\s+резьбовое\b', 'Калибр-кольцо', name, flags=re.IGNORECASE)
    name = re.sub(r'^Кольцо\b', 'Калибр-кольцо', name, flags=re.IGNORECASE)
    name = re.sub(r'^Пробка\s+резьбовая\b', 'Калибр-пробка', name, flags=re.IGNORECASE)
    name = re.sub(r'^Пробка\b', 'Калибр-пробка', name, flags=re.IGNORECASE)
    name = re.sub(r'^Скоба\b', 'Калибр-скоба', name, flags=re.IGNORECASE)
    # Символ ⌀ для скобы и пробки гладкой: «Калибр-скоба 1,2» → «Калибр-скоба ⌀1,2»
    name = re.sub(r'(Калибр-скоба|Калибр-пробка)\s+(?!⌀)(\d)', r'\1 ⌀\2', name, flags=re.IGNORECASE)
    # «Калибр М...» без типа — определяем по квалитету
    name = _infer_caliber_type(name)
    return ' '.join(name.split())


def _spec_part(name):
    """Вырезает 'М10×1,5 6g' из 'Калибр-кольцо М10×1,5 6g ПР'."""
    s = re.sub(r'^(Калибр-кольцо|Калибр-пробка|Кольцо|Пробка)\s*', '', name, flags=re.IGNORECASE)
    s = re.sub(r'\s+(ПР-НЕ|ПР|НЕ)\s*$', '', s, flags=re.IGNORECASE)
    return s.strip()


def _control_rows(ring_name, diam, pitch, ring_side, qty, include_kalib, discount):
    """
    Генерирует строки контрольных калибров-пробок для кольца.
    ring_side: 'ПР' → КПР-ПР, КПР-НЕ; 'НЕ' → КНЕ-ПР, КНЕ-НЕ
    """
    combo = _PLUG_COMBO.get((float(diam), float(pitch)))
    if not combo:
        return []
    tag = 'КПР' if ring_side == 'ПР' else 'КНЕ'
    spec = _spec_part(ring_name)
    d = discount if discount is not None else 0.0
    kal = 1760 if include_kalib else 0
    rows = []
    for ctrl_side in ('ПР', 'НЕ'):
        ctrl_name = f"Контрольный калибр-пробка {spec} {tag}-{ctrl_side}"
        disc    = round(combo * d, 2)
        price_d = round(combo - disc, 2)
        summa   = round(qty * (price_d + kal), 2)
        nds_s   = round(summa * NDS, 2)
        rows.append({
            'name': ctrl_name, 'ed': 'шт.', 'qty': qty,
            'price': combo, 'disc': disc, 'price_d': price_d,
            'kal': kal, 'summa': summa, 'nds': nds_s,
            'itogo': round(summa + nds_s, 2),
            'stock': None, 'stock_price': None,
            'source': 'control',
        })
    return rows


def _add_stock(row):
    """Добавляет информацию о складе в строку результата."""
    s = stock_lookup(row['name'])
    row['stock'] = s[1] if s else None
    row['stock_price'] = s[0] if s else None


def _needs_split(name):
    """
    Нужно ли разбивать позицию на ПР и НЕ.
    Кольца: всегда. Пробки резьбовые >50мм: да. Пробки ≤50мм: нет (цена комплекта).
    """
    nu = name.upper()
    if 'ПР-НЕ' not in nu and 'ПР/НЕ' not in nu:
        return False
    parsed = parse_caliber(name)
    if not parsed:
        return False
    kind = parsed[0]
    if kind == 'кольцо_резьбовое':
        return True
    if kind == 'пробка_резьбовая':
        diam = parsed[1]
        return float(diam) > 50
    return False


def _replace_pr_ne(name, side):
    """Заменяет ПР-НЕ / ПР/НЕ на ПР или НЕ, сохраняя регистр остальной части."""
    for combo in ('ПР-НЕ', 'ПР/НЕ'):
        idx = name.upper().find(combo)
        if idx >= 0:
            return name[:idx] + side + name[idx + len(combo):]
    return name


def _expand(name, qty, include_kalib, discount):
    """
    Считает одну или две позиции (при разбивке ПР/НЕ).
    Возвращает список dict-ов (1 или 2 элемента).
    """
    # Кол-во, вписанное в наименование, имеет приоритет
    qty_in_name = _extract_qty_from_name(name)
    if qty_in_name is not None:
        qty = qty_in_name

    if _needs_split(name):
        rows = []
        for side in ('ПР', 'НЕ'):
            n = _clean_name(_replace_pr_ne(name, side))
            r = calc_item(n, qty, include_kalib=include_kalib, discount=discount)
            if r:
                r['name'] = n
                _add_stock(r)
                rows.append(r)
                # Контрольные калибры для колец ≤36мм и шаг ≤1,5мм
                parsed = parse_caliber(n)
                if parsed and parsed[0] == 'кольцо_резьбовое':
                    d, p = parsed[1], parsed[2]
                    if float(d) <= 36 and float(p) <= 1.5:
                        rows.extend(_control_rows(n, d, p, side, qty,
                                                  include_kalib, discount))
        return rows

    clean = _clean_name(name)
    r = calc_item(clean, qty, include_kalib=include_kalib, discount=discount)
    if r:
        r['name'] = clean
        _add_stock(r)
        return [r]
    return []

# ─────────────────────────────────────────────────────────────────────────────
@app.route("/generate", methods=["POST"])
def generate():
    """Генерирует Word КП по данным формы."""
    if not session.get("auth"):
        return jsonify({"error": "not_auth"}), 403

    data = request.get_json(force=True)
    items     = data.get("items", [])
    aa        = [(x["name"], x["qty"]) for x in data.get("aa", [])]
    customer  = data.get("customer", "")
    order_num = data.get("order_num", "")
    manager   = data.get("manager", "Перог Вадим Александрович")
    try:
        discount = float(data.get("discount", 20)) / 100.0
        discount = max(0.0, min(0.99, discount))
    except (ValueError, TypeError):
        discount = 0.20

    if not items:
        return jsonify({"error": "Нет позиций для КП"}), 400

    buf = build_kp(items, aa, customer=customer,
                   order_num=order_num, manager=manager, discount=discount)

    fname = f"КП_{order_num}.docx" if order_num else "КП_ВЕГАЛИБР.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=fname,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# ─────────────────────────────────────────────────────────────────────────────
@app.route("/calc-single", methods=["POST"])
def calc_single():
    """Считает одну позицию по наименованию (ручное добавление)."""
    if not session.get("auth"):
        return jsonify({"error": "not_auth"}), 403

    data = request.get_json(force=True)
    name = (data.get("name") or "").strip()
    try:
        qty = max(1, int(data.get("qty", 1)))
    except (ValueError, TypeError):
        qty = 1
    include_kalib = bool(data.get("include_kalib", True))
    try:
        discount = float(data.get("discount", 20)) / 100.0
        discount = max(0.0, min(0.99, discount))
    except (ValueError, TypeError):
        discount = 0.20

    if not name:
        return jsonify({"error": "Укажите наименование"}), 400

    rows = _expand(name, qty, include_kalib, discount)
    if rows:
        return jsonify({"items": rows})

    return jsonify({"aa": {"name": name, "qty": qty}})


# ─────────────────────────────────────────────────────────────────────────────
@app.route("/parse-drawing", methods=["POST"])
def parse_drawing_route():
    """Принимает чертёж (PDF/PNG/JPG), возвращает список найденных калибров."""
    if not session.get("auth"):
        return jsonify({"error": "not_auth"}), 403

    f = request.files.get("file")
    if not f:
        return jsonify({"error": "Файл не загружен"}), 400

    file_bytes = f.read()
    filename   = f.filename or ""

    from drawing_parser import parse_drawing
    suggestions, err = parse_drawing(file_bytes, filename)
    if err and not suggestions:
        return jsonify({"error": err}), 400

    return jsonify({"suggestions": [{"name": n, "hint": h} for n, h in suggestions]})


# ─────────────────────────────────────────────────────────────────────────────
@app.route("/drawing-doc", methods=["POST"])
def drawing_doc():
    """Принимает чертёж → анализирует → возвращает Word-документ с запросом."""
    if not session.get("auth"):
        return jsonify({"error": "not_auth"}), 403

    f = request.files.get("file")
    if not f:
        return jsonify({"error": "Файл не загружен"}), 400

    file_bytes = f.read()
    filename   = f.filename or "чертёж"

    from vision_parser import analyze_drawing
    data = analyze_drawing(file_bytes, filename)

    if data['error'] and not data['items'] and not data['title']:
        return jsonify({"error": data['error']}), 400

    buf = _build_drawing_doc(
        title      = data['title']  or filename,
        number     = data['number'] or "",
        items      = data['items'],
        img_bytes  = file_bytes,
        img_name   = filename,
    )

    # Очищаем имя файла от символов запрещённых в Windows
    raw_title = data['title'] or "чертёж"
    safe_name = re.sub(r'[\\/:*?"<>|]', '', raw_title).strip().lstrip('_.- ')[:80] or "чертёж"
    dl_name   = f"{safe_name}.docx"

    import urllib.parse
    resp = send_file(
        buf,
        as_attachment=True,
        download_name=dl_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    # Дополнительный заголовок — JS читает его напрямую (обход проблем с кириллицей в Content-Disposition)
    resp.headers['X-Download-Filename'] = urllib.parse.quote(dl_name, safe='')
    return resp


def _build_drawing_doc(title: str, number: str, items,
                       img_bytes: bytes = None, img_name: str = "") -> BytesIO:
    """Собирает Word-документ с запросом на изготовление калибров."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin    = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.font.name  = "Arial"
            r.font.size  = Pt(size)
            if color:
                r.font.color.rgb = color
        return p

    # ── Шапка ─────────────────────────────────────────────────────────────────
    p = para("Уточнить у Александра Андреевича возможность изготовления",
             bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0xC0, 0x00, 0x00), space_before=0, space_after=16)

    # Горизонтальная линия
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), 'C00000')
    pb.append(bot); pPr.append(pb)

    # ── Реквизиты чертежа ──────────────────────────────────────────────────────
    para(f"Чертёж:  {title}", bold=True, size=12, space_before=10, space_after=4)
    if number:
        para(f"Обозначение:  {number}", size=11, space_before=0, space_after=10)
    else:
        para("", space_before=0, space_after=4)

    # ── Изображение чертежа ───────────────────────────────────────────────────
    if img_bytes:
        try:
            import io as _io
            fname_low = img_name.lower()
            img_buf = None

            if fname_low.endswith(('.jpg', '.jpeg', '.png')):
                img_buf = _io.BytesIO(img_bytes)
            elif fname_low.endswith('.pdf'):
                try:
                    import fitz as _fitz
                    _pdoc = _fitz.open(stream=img_bytes, filetype='pdf')
                    _pix  = _pdoc[0].get_pixmap(matrix=_fitz.Matrix(1.5, 1.5))
                    img_buf = _io.BytesIO(_pix.tobytes('png'))
                    _pdoc.close()
                except Exception:
                    img_buf = None
            else:   # tif, bmp и др.
                try:
                    from PIL import Image as _Img
                    _imo = _Img.open(_io.BytesIO(img_bytes)).convert("RGB")
                    img_buf = _io.BytesIO()
                    _imo.save(img_buf, format="PNG")
                    img_buf.seek(0)
                except Exception:
                    img_buf = None

            if img_buf is not None:
                # Вычисляем реальный размер изображения с учётом DPI
                # Не увеличиваем, если изображение меньше страницы
                kwargs = {"width": Cm(14)}   # безопасный fallback
                try:
                    from PIL import Image as _Pil
                    img_buf.seek(0)
                    _im2 = _Pil.open(img_buf)
                    w_px, h_px = _im2.size
                    _dpi = _im2.info.get('dpi', (150, 150))
                    dpi_x = float(_dpi[0]) if _dpi[0] and _dpi[0] > 0 else 150.0
                    dpi_y = float(_dpi[1]) if _dpi[1] and _dpi[1] > 0 else 150.0
                    # Физический размер изображения в см
                    nat_w = w_px / dpi_x * 2.54
                    nat_h = h_px / dpi_y * 2.54
                    # Ограничения страницы (не увеличиваем, только уменьшаем)
                    max_w, max_h = 17.0, 19.0
                    scale = min(1.0, max_w / nat_w, max_h / nat_h)
                    fw = nat_w * scale
                    fh = nat_h * scale
                    # Передаём ту сторону, которая задаёт масштаб
                    if fw / fh >= 1:
                        kwargs = {"width": Cm(fw)}
                    else:
                        kwargs = {"height": Cm(fh)}
                except Exception:
                    pass   # используем fallback

                p_img = doc.add_paragraph()
                p_img.paragraph_format.space_before = Pt(6)
                p_img.paragraph_format.space_after  = Pt(6)
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                img_buf.seek(0)
                run_img.add_picture(img_buf, **kwargs)
        except Exception:
            pass   # если не удалось — просто пропускаем картинку

    # ── Список калибров ────────────────────────────────────────────────────────
    para("Необходимые калибры:", bold=True, size=11, space_before=4, space_after=6)

    if items:
        for i, (name, hint) in enumerate(items, 1):
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after  = Pt(2)
            p2.paragraph_format.left_indent  = Cm(0.5)
            r1 = p2.add_run(f"{i}.  {name}")
            r1.font.name = "Arial"; r1.font.size = Pt(11)
            if hint:
                r2 = p2.add_run(f"   ({hint})")
                r2.font.name = "Arial"; r2.font.size = Pt(9)
                r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    else:
        para("(калибры не определены — требуется ручной анализ)",
             size=10, color=RGBColor(0x88, 0x88, 0x88), space_before=0, space_after=6)

    # ── Подпись ────────────────────────────────────────────────────────────────
    para("", space_before=16, space_after=0)
    from datetime import date
    para(f"Дата:  {date.today().strftime('%d.%m.%Y')}",
         size=10, color=RGBColor(0x88, 0x88, 0x88), space_before=0, space_after=0)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
