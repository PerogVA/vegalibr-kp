# -*- coding: utf-8 -*-
"""Генератор Word КП для ВЕГАЛИБР"""

import os
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pricing import NDS, DISC

# Пути к картинкам (относительно этого файла)
_HERE = os.path.dirname(os.path.abspath(__file__))
LOGO_PNG    = os.path.join(_HERE, "static", "logo.png")
PRODUCT_PNG = os.path.join(_HERE, "static", "cat1.jpg")

# ── Цвета ────────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1F, 0x5C, 0x9E)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF7)
BG_BLUE    = RGBColor(0xEE, 0xF4, 0xFB)
STRIPE     = RGBColor(0xF5, 0xF9, 0xFF)
DARK_GRAY  = RGBColor(0x40, 0x40, 0x40)
MID_GRAY   = RGBColor(0xAA, 0xAA, 0xAA)
DARK_BG    = RGBColor(0x1A, 0x1A, 0x2E)

def _hex(c): return f"{c[0]:02X}{c[1]:02X}{c[2]:02X}"
def _fmt(v): return f"{v:,.2f}".replace(",", " ").replace(".", ",")

def _set_bg(cell, color):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), _hex(color)); pr.append(s)

def _set_border(cell, **kw):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for side, cfg in kw.items():
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"),   cfg.get("val","single"))
        e.set(qn("w:sz"),    str(cfg.get("sz", 4)))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), cfg.get("color", "000000"))
        b.append(e)
    pr.append(b)

def _no_border(cell):
    _set_border(cell, **{s: {"val":"nil","sz":0,"color":"FFFFFF"}
                         for s in ("top","bottom","left","right")})

def _margin(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for n, v in (("top",top),("bottom",bottom),("left",left),("right",right)):
        e = OxmlElement(f"w:{n}"); e.set(qn("w:w"), str(v)); e.set(qn("w:type"), "dxa")
        m.append(e)
    pr.append(m)

def _run(para, text, bold=False, size=10, color=None):
    r = para.add_run(text); r.bold = bold
    r.font.name = "Arial"; r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

def _cell_para(cell, text="", bold=False, size=10, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.alignment = align
    if text: _run(p, text, bold=bold, size=size, color=color)
    return p

def _add_para(cell, text="", bold=False, size=10, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=0):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.alignment = align
    if text: _run(p, text, bold=bold, size=size, color=color)
    return p

def _no_tbl_borders(table):
    tbl = table._tbl; pr = tbl.find(qn("w:tblPr"))
    if pr is None: pr = OxmlElement("w:tblPr"); tbl.insert(0, pr)
    b = OxmlElement("w:tblBorders")
    for s in ("top","bottom","left","right","insideH","insideV"):
        e = OxmlElement(f"w:{s}"); e.set(qn("w:val"),"none"); e.set(qn("w:sz"),"0")
        e.set(qn("w:space"),"0"); e.set(qn("w:color"),"FFFFFF"); b.append(e)
    pr.append(b)

def _thin(color="CCCCCC"): return {"val":"single","sz":4,"color":color}
def _thick(color, sz=8):   return {"val":"single","sz":sz,"color":color}
def _merge(row, a, b):
    c = row.cells[a]; c.merge(row.cells[b]); return c


def build_kp(items, aa_list, customer="", order_num="", manager="Перог Вадим Александрович",
             discount=0.20):
    """
    items: list of dict (from pricing.calc_item)
    aa_list: list of (name, qty) — позиции без цены
    Возвращает BytesIO с .docx
    """
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0);  sec.right_margin = Cm(1.0)
    sec.top_margin  = Cm(1.0);  sec.bottom_margin = Cm(1.0)
    style = doc.styles["Normal"]
    style.font.name = "Arial"; style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)

    # ── Шапка ────────────────────────────────────────────────────────────────
    th = doc.add_table(rows=1, cols=3)
    th.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_tbl_borders(th)

    pc = th.cell(0,0); pc.width = Cm(5.5)
    _set_bg(pc, RGBColor(0xF0,0xF0,0xF0))
    _margin(pc, 0, 0, 0, 0); _no_border(pc)
    pc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if os.path.exists(PRODUCT_PNG):
        _cell_para(pc,"",align=WD_ALIGN_PARAGRAPH.CENTER).add_run().add_picture(PRODUCT_PNG,width=Cm(5.5))

    lc = th.cell(0,1); lc.width = Cm(7.5)
    _set_bg(lc, DARK_BG); _margin(lc,60,60,100,0); _no_border(lc)
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if os.path.exists(LOGO_PNG):
        _cell_para(lc,"",align=WD_ALIGN_PARAGRAPH.LEFT).add_run().add_picture(LOGO_PNG,width=Cm(7.2),height=Cm(1.28))

    ic = th.cell(0,2); ic.width = Cm(5.0)
    _set_bg(ic, DARK_BG); _margin(ic,60,60,140,100); _no_border(ic)
    ic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_para(ic,"ООО «ВЕГАЛИБР»",bold=True,size=9.5,color=WHITE,sa=3)
    _add_para(ic,"www.vegalibr.ru",size=8,color=RGBColor(0xAA,0xCC,0xFF),sa=2)
    _add_para(ic,"ИНН 9731016281",size=7.5,color=RGBColor(0xAA,0xAA,0xCC),sa=3)
    _add_para(ic,"+7 (936) 617 61 63",size=8.5,color=WHITE,sa=1)
    _add_para(ic,"sales3@vegalibr.ru",size=8,color=WHITE)

    # ── Разделитель ──────────────────────────────────────────────────────────
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(3)
    div.paragraph_format.space_after  = Pt(3)
    pPr = div._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"12")
    bot.set(qn("w:space"),"1");    bot.set(qn("w:color"),_hex(BLUE))
    pBdr.append(bot); pPr.append(pBdr)

    # ── Заголовок ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(1)
    _run(p, "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", bold=True, size=13, color=BLUE)

    order_label = f"Заявка {order_num}  |  от ___.___._______" if order_num else "от ___.___._______"
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    _run(p2, order_label, size=9.5, color=MID_GRAY)

    # ── Заказчик ──────────────────────────────────────────────────────────────
    def _label_row(tbl, label, value=""):
        row = tbl.add_row(); lc2, vc = row.cells[0], row.cells[1]
        lc2.width = Cm(4.2); vc.width = Cm(13.8)
        _no_border(lc2); _no_border(vc)
        _margin(lc2,40,40,0,80); _margin(vc,40,40,60,0)
        _cell_para(lc2, label, bold=True, size=9, color=DARK_GRAY)
        pv = _cell_para(vc, value or " ", size=9)
        ppPr = pv._p.get_or_add_pPr(); pb = OxmlElement("w:pBdr")
        b2 = OxmlElement("w:bottom")
        b2.set(qn("w:val"),"single"); b2.set(qn("w:sz"),"4")
        b2.set(qn("w:space"),"1");    b2.set(qn("w:color"),"CCCCCC")
        pb.append(b2); ppPr.append(pb)

    tcl = doc.add_table(rows=0, cols=2)
    tcl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_tbl_borders(tcl)
    _label_row(tcl, "Заказчик:", customer)
    _label_row(tcl, "Контактное лицо:")

    # ── Таблица позиций ───────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(3)

    HEADERS = ["№","Наименование","Ед.","Кол-во","Цена\nбез НДС","Скидка\n20%",
               "Цена со\nскидкой","Кали-\nбровка","Сумма\nбез НДС","НДС\n22%","Итого\nс НДС"]
    COL_W   = [0.6, 5.5, 0.65, 0.8, 1.5, 1.3, 1.5, 1.3, 1.5, 1.25, 1.5]

    tbl = doc.add_table(rows=1, cols=11)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Table Grid"

    hdr = tbl.rows[0]; hdr.height = Cm(0.75)
    for cell, text, w in zip(hdr.cells, HEADERS, COL_W):
        cell.width = Cm(w)
        _set_bg(cell, BLUE)
        _set_border(cell, top=_thin(), bottom=_thin(), left=_thin(), right=_thin())
        _margin(cell, 60, 60, 40, 40)
        _cell_para(cell, text, bold=True, size=7.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total_bez  = 0.0
    total_disc = 0.0
    total_kal  = 0.0

    for idx, it in enumerate(items):
        total_bez  += it['summa']
        total_disc += it['qty'] * it['disc']
        total_kal  += it['qty'] * it['kal']

        row = tbl.add_row(); row.height = Cm(0.55)
        bg = STRIPE if idx % 2 == 1 else None
        vals = [str(idx+1), it['name'], it['ed'], str(it['qty']),
                _fmt(it['price']), _fmt(it['disc']), _fmt(it['price_d']),
                _fmt(it['kal']),   _fmt(it['summa']), _fmt(it['nds']), _fmt(it['itogo'])]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.RIGHT,  WD_ALIGN_PARAGRAPH.RIGHT,
                  WD_ALIGN_PARAGRAPH.RIGHT]
        for cell, v, w, al in zip(row.cells, vals, COL_W, aligns):
            cell.width = Cm(w)
            if bg: _set_bg(cell, bg)
            _set_border(cell, top=_thin(), bottom=_thin(), left=_thin(), right=_thin())
            _margin(cell, 40, 40, 40, 40)
            _cell_para(cell, v, size=7.5, color=DARK_GRAY, align=al)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    total_bez  = round(total_bez, 2)
    total_disc = round(total_disc, 2)
    total_nds  = round(total_bez * NDS, 2)
    total_tot  = round(total_bez + total_nds, 2)

    def _summary(label, val, bg=LIGHT_BLUE, color=BLUE, sz=8.5):
        r = tbl.add_row()
        mc = _merge(r, 0, 7)
        _set_bg(mc, bg); _set_border(mc, top=_thin(), bottom=_thin(), left=_thin(), right=_thin())
        _margin(mc, 60, 60, 80, 80)
        _cell_para(mc, label, bold=True, size=sz, color=color, align=WD_ALIGN_PARAGRAPH.RIGHT)
        mc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        mv = _merge(r, 8, 10)
        _set_bg(mv, bg); _set_border(mv, top=_thin(), bottom=_thin(), left=_thin(), right=_thin())
        _margin(mv, 60, 60, 60, 60)
        _cell_para(mv, val, bold=True, size=sz, color=color, align=WD_ALIGN_PARAGRAPH.RIGHT)
        mv.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _summary("в т.ч. скидка 20% на изделия:", f"– {_fmt(total_disc)}",
             bg=RGBColor(0xFF,0xF0,0xD0), color=RGBColor(0xCC,0x44,0x00))
    _summary("в т.ч. калибровка:", _fmt(round(total_kal,2)))
    _summary("Итого без НДС:", _fmt(total_bez))
    _summary("НДС 22%:", _fmt(total_nds))

    r_tot = tbl.add_row()
    mc2 = _merge(r_tot, 0, 7)
    _set_bg(mc2, LIGHT_BLUE)
    _set_border(mc2, top=_thin(), bottom=_thick(_hex(BLUE),10), left=_thin(), right=_thin())
    _margin(mc2, 80, 80, 80, 80)
    _cell_para(mc2, "ИТОГО с НДС:", bold=True, size=10, color=BLUE, align=WD_ALIGN_PARAGRAPH.RIGHT)
    mc2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    mc3 = _merge(r_tot, 8, 10)
    _set_bg(mc3, LIGHT_BLUE)
    _set_border(mc3, top=_thin(), bottom=_thick(_hex(BLUE),10), left=_thin(), right=_thin())
    _margin(mc3, 80, 80, 60, 60)
    _cell_para(mc3, _fmt(total_tot) + " руб.", bold=True, size=10, color=BLUE,
               align=WD_ALIGN_PARAGRAPH.RIGHT)
    mc3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Условия ──────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(4)
    tc = doc.add_table(rows=1, cols=1); tc.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_tbl_borders(tc)
    cc = tc.cell(0,0)
    _set_bg(cc, BG_BLUE)
    _set_border(cc, top=_thick(_hex(BLUE),8),
                bottom={"val":"nil","sz":0,"color":"FFFFFF"},
                left=_thick(_hex(BLUE),16),
                right={"val":"nil","sz":0,"color":"FFFFFF"})
    _margin(cc, 60, 60, 140, 120)
    _cell_para(cc, "Условия поставки:", bold=True, size=8.5, color=BLUE, sa=2)
    for line in [
        "Срок изготовления: 50–70 рабочих дней  |  Ускоренно (до 30 дней): +100%  |  Действие КП: 10 раб. дней",
        "Оплата: 100% предоплата  |  Доставка: СДЭК / Почта России / самовывоз  |  Цены без НДС (НДС 22%)",
    ]:
        _add_para(cc, line, size=8, color=DARK_GRAY, sa=1)

    # ── Уточнить у АА ────────────────────────────────────────────────────────
    if aa_list:
        doc.add_paragraph().paragraph_format.space_before = Pt(3)
        p_aa = doc.add_paragraph(); p_aa.paragraph_format.space_before = Pt(2)
        _run(p_aa, "Требуют уточнения у АА (размер <1,0 мм или нет в прайсе):",
             bold=True, size=8, color=RGBColor(0xCC,0x00,0x00))
        for name, qty in aa_list:
            pa = doc.add_paragraph(); pa.paragraph_format.space_before = Pt(1)
            _run(pa, f"  — {name}  ×{qty} шт.", size=8, color=DARK_GRAY)

    # ── Подпись ──────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_before = Pt(5)
    ts = doc.add_table(rows=1, cols=3); ts.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_tbl_borders(ts)
    sl, sn, ss = ts.cell(0,0), ts.cell(0,1), ts.cell(0,2)
    sl.width = Cm(5.0); sn.width = Cm(7.0); ss.width = Cm(6.0)
    for c in (sl, sn, ss): _no_border(c)
    _margin(sl, 40,40, 0,60); _margin(sn, 40,40, 60,60); _margin(ss, 40,40, 60,0)
    _cell_para(sl, "Менеджер по продажам:", size=8.5, color=DARK_GRAY)
    _cell_para(sn, manager, bold=True, size=8.5, color=DARK_GRAY)
    _cell_para(ss, "___________________", size=8.5, color=MID_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── Сохраняем в память ───────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
